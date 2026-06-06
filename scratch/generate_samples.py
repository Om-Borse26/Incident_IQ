
from pathlib import Path
import docx
from docx.shared import Pt
import fitz # PyMuPDF

# Read MD
md_path = Path("d:/02_Interests (Up Skilling )/05_LLM, RAG, Agents, MCP and more/Project - Incident_IQ/data/incidents/notification-service-queue-backlog.md")
text = md_path.read_text(encoding="utf-8")

# --- Create DOCX ---
doc = docx.Document()
lines = text.split("\n")
for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=2)
    else:
        doc.add_paragraph(line)

docx_path = md_path.with_suffix(".docx")
doc.save(str(docx_path))
print(f"Saved {docx_path}")

# --- Create PDF ---
# We'll use fitz.Document and insert text
pdf_doc = fitz.open()
page = pdf_doc.new_page()
y = 50
for line in lines:
    line = line.strip()
    if not line:
        continue
    if line.startswith("# "):
        page.insert_text((50, y), line[2:], fontsize=18)
        y += 24
    elif line.startswith("## "):
        y += 10
        page.insert_text((50, y), line[3:], fontsize=14)
        y += 18
    else:
        # crude wrapping
        if len(line) > 80:
            page.insert_text((50, y), line[:80], fontsize=10)
            y += 12
            page.insert_text((50, y), line[80:], fontsize=10)
        else:
            page.insert_text((50, y), line, fontsize=10)
        y += 12

pdf_path = md_path.with_suffix(".pdf")
pdf_doc.save(str(pdf_path))
pdf_doc.close()
print(f"Saved {pdf_path}")
